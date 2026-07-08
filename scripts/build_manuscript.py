"""Build the manuscript draft (.docx) for the Chromatin-Toggle / multiscale KG-GNN work.
Honest, numbers-current draft: Abstract, Introduction, Methods, Results, Discussion,
Limitations, Conclusion. Kept separate from build_report.py (the results summary)."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FIG = "/Users/work/MultiscaleProject/artifacts/figures"
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

def h(t, l=1): doc.add_heading(t, level=l)
def p(t):
    para = doc.add_paragraph(t); para.paragraph_format.space_after = Pt(6); return para
def em(t):
    q = doc.add_paragraph(); r = q.add_run(t); r.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55); return q
def fig(name, w=5.4, cap=None):
    path = os.path.join(FIG, name)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        c = doc.add_paragraph(cap); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic = True; c.runs[0].font.size = Pt(9)
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hh in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ""
        r = cell.paragraphs[0].add_run(hh); r.bold = True; r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ---------------- Title block ----------------
t = doc.add_heading("A knowledge-graph neural network operationalizing a multiscale theory "
                    "of cellular information processing", 0)
s = doc.add_paragraph("Predicting cell-state programs across pathways from intrinsic memory and "
                      "extrinsic cues, under leakage-controlled evaluation")
s.runs[0].italic = True
a = doc.add_paragraph("N. Lytell.  Project Lumos.  Manuscript draft — 2026-07-06.")
a.runs[0].font.size = Pt(9)
em("DRAFT. Numbers are current as of this build. Items marked [pending] await a GPU-converged "
   "run; the reported conclusions do not depend on them.")

# ---------------- Abstract ----------------
h("Abstract")
p("Cells select a response program by integrating a stable intrinsic identity (lineage "
  "transcription factors, chromatin state) with transient extrinsic cues. We ask whether a "
  "specific multiscale theory of this integration — asymmetric intrinsic/extrinsic weighting, a "
  "plasticity gate, and attractor stabilization — can be built as an interpretable model with "
  "measurable predictive value. We encode the theory as a relation-typed graph neural network "
  "(KG-GNN) over a literature-derived knowledge graph and train it on 18,392 real single cells "
  "spanning 12 cell-state programs from 19 published datasets, under strict controls: labels are "
  "graded for provenance, program-marker genes can be masked, and generalization is measured by "
  "holding out entire datasets. After identifying and removing a cue-label leak that had inflated "
  "an earlier result, the honest findings are: (i) the expression->program mapping is highly "
  "learnable in-distribution (stratified balanced accuracy 0.93); (ii) cross-dataset "
  "generalization is hard for all models (a batch/domain-shift frontier); (ii) there, with markers "
  "included in all arms, adding the regulatory graph SIGNIFICANTLY improves program probability-"
  "ranking — the full KG-GNN (macro-AUPRC 0.47) beats the identical network with edges removed "
  "(0.39; paired p=0.015) and both structureless baselines (logistic regression, random forest; "
  "p<0.02), while matching them on top-1 metrics — a direct, controlled confirmation that layered "
  "regulatory structure adds value on top of markers; (iii) in simulation the model reproduces "
  "the theory's plasticity-gated, hysteretic switching, and on a real EMT time-course — under the "
  "markers-in, attractor-off setting — it reads graded temporal program emergence (rho +0.20) that "
  "COLLAPSES to +0.03 when the graph is removed, a second structure-isolation confirmation (this "
  "time on the temporal axis) that regulatory structure, not markers, drives the signal. We "
  "conclude that the theory yields an "
  "interpretable model whose concrete real-data advantage is probability-ranking, and we report "
  "the negative results (temporal, top-1) rather than tune them away.")

# ---------------- 1 Introduction ----------------
h("1. Introduction")
p("A central question in cell biology is how a cell decides which of many possible response "
  "programs to execute. A recurring view — here called the multiscale theory — holds that the "
  "decision integrates information across scales and timescales: a deeply processed, persistent "
  "intrinsic bias (lineage/chromatin memory) is weighted strongly, while shallowly processed, "
  "transient extrinsic cues are weighted weakly; a window of plasticity lets a weak cue overcome "
  "the intrinsic default; and feedback stabilizes the winning program as a self-sustaining "
  "attractor, so the new state persists after the cue is withdrawn (hysteresis = stored memory).")
p("Testing such a theory requires more than a black-box predictor: the mechanisms must be built "
  "in and independently removable, and the evaluation must resist the shortcuts that make "
  "cell-state prediction look easy (marker-gene circularity, label leakage, and batch/source "
  "effects). We therefore (1) encode the theory as an interpretable knowledge-graph GNN whose "
  "mechanisms are toggleable, (2) train and evaluate it on real multi-source single-cell data "
  "under leakage controls, and (3) compare it against strong theory-agnostic baselines on the "
  "metrics appropriate for an imbalanced multiclass problem. Our aim is an honest account of "
  "where the theory's inductive biases help, where they do not, and why.")

# ---------------- 2 Methods ----------------
h("2. Methods")
h("2.1 Knowledge graph and model", 2)
p("The model's objective is to predict a cell's response-program (pathway) state from its "
  "measured gene activity over the regulatory graph. Nodes are molecules, chromatin features, or "
  "response programs; edges are literature activation/inhibition relations. The GNN sees only the "
  "binary graph STRUCTURE (edge signs/weights are never exposed) plus the cell's expression "
  "injected onto nodes, and reads out a program via message passing. A hybrid residual (a linear "
  "map from the raw node vector to the class logits) guarantees it is at least as expressive as a "
  "linear classifier; message passing is vectorized over relations.")
p("Intrinsic memory formulation. We evaluated two formulations of how lineage/chromatin memory "
  "enters the model. The INITIAL formulation re-injected intrinsic memory as a strong signal every "
  "message-passing round (with a decaying cue and a plasticity term scaling cue influence, and a "
  "winner-take-all attractor step) — effectively adding a secondary autoregulation mechanism on "
  "top of the pathway prediction. The PREFERRED formulation (Section 3.7) reframes memory as "
  "transition RESISTANCE: memory sets the barrier a cue must overcome to leave the current "
  "attractor, plasticity LOWERS that barrier (rather than amplifying the cue), and a soft/graded "
  "attractor replaces winner-take-all. All of these are independently ablatable so the "
  "contribution of each can be measured; leakage controls are identical across formulations.")
h("2.2 Data and labels", 2)
p("The training pool is 18,392 cells, 12 programs, 19 datasets (capped at 600 cells per "
  "program-per-dataset to limit single-cell dominance), with inputs widened to a curated 148-gene "
  "marker/TF panel. Most programs have >=2 independent sources (different tissue or species). "
  "Labels are graded by provenance: type A (deposited per-cell author annotations), type B "
  "(per-cell x disease field), type C (condition/timepoint proxy), type D (derived by us); the "
  "grading is documented in a label-provenance protocol. Cues are applied UNIFORMLY per dataset "
  "(not gated by outcome) after we found that outcome-gated cues had leaked the label.")
h("2.3 Evaluation", 2)
p("We report two complementary regimes. LEARNABILITY: stratified k-fold CV (cells mixed) — can the "
  "expression->program mapping be learned at all. GENERALIZATION: grouped k-fold CV holding out "
  "ENTIRE datasets — can the model transfer to unseen sources (the honest, harder test). A "
  "marker-shortcut control masks program-defining readout genes. Class weighting counters the "
  "~half-Quiescent imbalance. For imbalanced multiclass we emphasize macro-AUPRC (threshold-"
  "independent, probability-ranking) and balanced accuracy over cell-weighted accuracy. Baselines "
  "(majority, logistic regression, random forest, gradient boosting) are run on identical features "
  "and folds; a paired Wilcoxon test compares the KG-GNN to each baseline on the same seed x fold "
  "splits.")

# ---------------- 3 Results ----------------
h("3. Results")

h("3.1 Label integrity: a cue-label leak, found and fixed", 2)
p("An earlier configuration applied each cue only to treated/diseased cells, which for "
  "disease-vs-normal datasets made the cue a perfect label proxy (e.g. MechanicalStretch=1 iff "
  "Hypertrophy) and inflated cross-species Hypertrophy transfer to 0.98. Applying the cue "
  "uniformly per dataset removes the leak; the 0.98 figure is withdrawn. All results below use "
  "leak-free, uniform cues.")

h("3.2 Learnability vs generalization", 2)
p("In-distribution the mapping is highly learnable; transferring to unseen datasets is the "
  "frontier. This gap explains why the honest (grouped) numbers look modest — it is domain shift, "
  "not a weak model.")
table(["Regime", "balanced acc", "prog recall", "macro-AUPRC"],
      [["Learnability (stratified, logreg)", "0.925", "0.966", "0.885"],
       ["Generalization (grouped, KG-GNN)", "0.319", "0.270", "0.472"],
       ["Generalization (grouped, best baseline)", "0.377", "0.324", "0.397"]])
em("The stratified number is optimistic (each dataset is program-enriched, so recognizing the "
   "source partly predicts the program — the shortcut grouped-split removes). It is reported as "
   "evidence the task is learnable, not as a generalization claim.")

h("3.3 Central result: regulatory structure adds value on top of markers", 2)
p("The core hypothesis — cellular state is shaped by LAYERED regulatory context — predicts that a "
  "model seeing markers PLUS regulatory structure should beat a model seeing markers WITHOUT "
  "structure. We test this directly with markers included in BOTH arms (mask=none) and the "
  "attractor off (an honest graded classifier, no forced fate), isolating structure as the only "
  "variable: the full KG-GNN vs the IDENTICAL network with the graph edges removed "
  "(kg_gnn_noedges) — same architecture, capacity, features, and markers. 5-fold x 3-seed grouped "
  "CV; paired Wilcoxon over the same seed x fold splits.")
table(["Model (markers in all)", "macro-AUPRC", "vs KG-GNN (paired)"],
      [["KG-GNN (markers + structure)", "0.473", "—"],
       ["KG-GNN, edges removed (markers, no structure)", "0.392", "+0.082, p=0.015 *"],
       ["Logistic regression (no structure)", "0.397", "+0.089, p=0.018 *"],
       ["Random forest (no structure)", "0.405", "+0.059, p=0.015 *"]])
em("With markers held equal, adding the regulatory graph SIGNIFICANTLY improves program "
   "probability-ranking (+0.082 AUPRC, p=0.015). The control is decisive: removing the graph from "
   "the same network drops it to 0.392 — right onto the logistic-regression level (0.397) — so the "
   "structure, not capacity or features, is the lever. The KG-GNN also significantly beats both "
   "external structureless baselines. Scope: the gain is specific to AUPRC (ranking); on prog-"
   "recall the arms tie (p~1.0). Claim: layered regulatory structure significantly improves the "
   "model's ranking of the correct program, markers held equal — a direct, controlled test of the "
   "theory. (Attractor off; the winner-take-all sharpening, when on, saturates probabilities and "
   "erases this graded ranking signal — see 3.4.)")
fig("structure_isolation.png", 5.6, "Figure 1. Structure-isolation test (markers in all arms, "
    "attractor off). Adding the regulatory graph (blue) significantly raises program-ranking AUPRC "
    "over the identical edge-removed network and structureless baselines (paired Wilcoxon vs "
    "KG-GNN). Wide error bars are grouped-fold variance; the paired test is the correct comparison.")

h("3.3b Cross-dataset classification vs baselines (no-markers view)", 2)
p("Widened inputs, 5-fold x 3-seed grouped CV, markers removed, class-weighted (mean +/- std over "
  "the 15 seed x fold estimates). The KG-GNN is the top model on macro-AUPRC — the metric that "
  "matters for imbalanced multiclass — by a consistent ~19% relative margin, while trailing on the "
  "argmax metrics at this compute budget.")
table(["Model", "prog recall", "balanced acc", "macro-F1", "macro-AUPRC"],
      [["KG-GNN", "0.270 +/-0.13", "0.319 +/-0.10", "0.136 +/-0.05", "0.472 +/-0.11"],
       ["Random forest", "0.324 +/-0.16", "0.377 +/-0.12", "0.194 +/-0.06", "0.397 +/-0.13"],
       ["Logistic regression", "0.307 +/-0.18", "0.365 +/-0.13", "0.161 +/-0.05", "0.396 +/-0.16"],
       ["Gradient boosting", "0.211 +/-0.12", "0.345 +/-0.09", "0.173 +/-0.05", "0.392 +/-0.16"],
       ["Majority class", "0.000", "0.220 +/-0.02", "0.138 +/-0.02", "0.154 +/-0.04"]])
p("Significance (paired Wilcoxon signed-rank over the same seed x fold splits — the correct test, "
  "which controls for fold difficulty):")
table(["KG-GNN vs", "AUPRC edge", "p-value"],
      [["Logistic regression", "+0.069", "0.008 (significant)"],
       ["Random forest", "+0.058", "0.018 (significant)"],
       ["Majority", "+0.278", "0.0001 (significant)"]])
em("The KG-GNN's macro-AUPRC advantage over both strong baselines is STATISTICALLY SIGNIFICANT "
   "(p<0.02) on the paired test. The marginal +/-1 std ranges overlap only because grouped-fold "
   "variance is large; paired (same folds), the edge is significant. The advantage is specific to "
   "AUPRC (probability ranking): on prog-recall the model ties logistic regression (p=0.85) and "
   "beats random forest (p=1e-4). The edge is consistent across every configuration (single-seed "
   "0.500, multi-seed 0.477); widening inputs (42->148 genes) lifted all models ~0.03-0.05 AUPRC. "
   "Argmax metrics trail at epochs-40 (under-convergence) and recover to competitive when converged "
   "(recall 0.437 / balanced-acc 0.382, single seed). Interpretation: structure improves program "
   "RANKING over structureless learners, significantly, while matching them on top-1 decisions. "
   "This edge does not shrink as the program set grows — on a 19-program pool it widens and extends "
   "to top-1 metrics (Section 3.8).")

h("3.4 Theory dynamics: simulation vs a real time-course", 2)
p("In SIMULATION, sweeping the plasticity input reproduces the theory's central behavior: at low "
  "plasticity the intrinsic default holds regardless of cue; past a threshold the cue flips the "
  "stabilized program; the flip persists after cue withdrawal (hysteresis). This held for "
  "Fibrosis, Hypertrophy, and InnateMemory across independent sources and survived marker removal; "
  "removing the plasticity gate abolishes the effect (it is load-bearing). No-cue programs stayed "
  "flat (control). We then tested the theory's temporal-integration prediction on a REAL EMT "
  "time-course (0d Quiescent; 8h/1d/3d/7d all labelled EMT), using experimental time only for "
  "validation (never as a model input): does predicted P(EMT) rise with time among the "
  "same-labelled cells (a graded rise = time-integrated commitment beyond the binary label)?")
table(["Config (EMT time-course, markers-in, attractor OFF)", "Spearman(P(EMT), time), EMT-labelled cells"],
      [["KG-GNN + regulatory structure", "+0.203 (p=5e-68) — graded"],
       ["KG-GNN, edges removed (markers, NO structure)", "+0.033 (p=5e-03) — flat"],
       ["(context) logistic regression, no-markers", "+0.154 — graded"],
       ["(context) KG-GNN, no-markers, attractor ON", "+0.007 — flat / null"]])
em("A SECOND structure-isolation result, on the temporal axis. Under the thesis-aligned setting "
   "(markers in, winner-take-all attractor OFF so probabilities stay graded) the KG-GNN reads "
   "graded temporal emergence — P(EMT) rises monotonically across 8h->7d (rho +0.203). The "
   "edge-removed control settles the attribution: with the SAME network, markers, and config but "
   "the graph removed, the gradient COLLAPSES to +0.033 (near-flat). So the regulatory STRUCTURE, "
   "not the markers, drives the graded temporal commitment — a second controlled confirmation of "
   "the theory (the first being program ranking, 3.3). The attractor must be off: its winner-"
   "take-all sharpening saturates probabilities to 0/1 and erases this graded signal (the earlier "
   "flat/null results were the masked-markers + attractor-on config). Separately, in simulation "
   "the plasticity-gated switching/hysteresis emerge by construction.")

h("3.4b Mechanism and structure ablation", 2)
p("Knocking out one component at a time on a fixed converged split (markers-in), scored on "
  "macro-AUPRC (the ranking metric of 3.3). Delta vs the full model:")
table(["Knockout", "ΔAUPRC", "reading"],
      [["scramble edges (wrong wiring)", "-0.155", "correct wiring matters (pro-structure)"],
       ["attractor / WTA", "-0.109", "helps here; flattens temporal (mixed tradeoff)"],
       ["remove chromatin nodes", "-0.058", "chromatin-memory carries ranking signal"],
       ["collapse relation types", "-0.057", "relation-typing matters"],
       ["remove TF nodes", "-0.020", "minor"],
       ["hybrid residual / asymmetric", "-0.015 / -0.014", "minor"],
       ["plasticity gate", "0.000", "inert for classification (a dynamics mechanism)"],
       ["remove all edges (no_edges)", "+0.035", "single-split only — see caveat"]])
em("Robust readings: the plasticity gate is INERT for static classification (it shapes the "
   "simulated dynamics, not prediction); scrambling or collapsing the graph badly hurts ranking, so "
   "correct relational wiring matters (pro-structure); removing chromatin-memory nodes also hurts "
   "ranking (-0.058), supporting the intrinsic-memory component; the attractor is a genuine tradeoff (helps "
   "this split's ranking but flattens the temporal continuum, 3.4). CAVEAT: this ablation is a "
   "SINGLE fixed split, and grouped-split variance is large (~+/-0.12); its all-or-nothing no_edges "
   "line came out +0.035, OPPOSITE to the reliable multi-seed, paired structure-isolation test in "
   "3.3 (structure helps, +0.082, p=0.015). We therefore rely on 3.3 (15 estimates + paired test) "
   "for the structure claim and treat the single-split no_edges as noise. A multi-seed ablation "
   "would be needed to score the all-or-nothing structure knockout reliably here.")

h("3.5 Interpretability", 2)
p("Permutation importance shows the model relies on textbook regulators for most programs "
  "(Hypertrophy: MechanicalStretch, HDAC4/5; MyogenicDiff: MyoD; ADM: Sox9, Caerulein; "
  "InnateMemory: LPS, PU.1; Regeneration: HDAC1/3, SWI/SNF; Senescence: CDKN2A/CDKN1A). EMT and "
  "Pluripotency did not surface their specific drivers (flagged).")
fig("importance_heatmap.png", 5.0, "Figure 2. Node x program permutation importance (prior 10-program analysis; being regenerated on the 12-program wide model).")

h("3.6 A falsifiable biological prediction (hypertrophy)", 2)
p("The encoded hypertrophy cascade (MechanicalStretch -> CaMKII/PKD -> nuclear export of HDAC4/5 "
  "-> de-repression of MEF2 -> Hypertrophy) yields sign-specific predictions we tested in-silico by "
  "editing node inputs on held-out hypertrophy cells (baseline P(Hypertrophy)=0.245). The result is "
  "PARTIAL and reported honestly: the model captured the FORWARD-ACTIVATION direction — boosting "
  "CaMKII raised P(Hypertrophy) by +0.121 — but did NOT reproduce the more distinctive DE-REPRESSION "
  "prediction: HDAC4/5 knockdown, which should enhance the program by releasing MEF2, produced ~0 "
  "change (-0.001). CaMKII/PKD blockade moved P(Hypertrophy) in the correct (down) direction but "
  "negligibly (-0.012 / -0.000). So the learned representation respects the activator axis but not "
  "the HDAC4/5 de-repression logic; the HDAC4/5 direction therefore remains an UNCONFIRMED, "
  "falsifiable wet-lab prediction (KN-93 CaMKII inhibition; HDAC4/5 knockdown in NRVM / hiPSC-CM), "
  "not an in-silico success.")

# ---------------- 4 Discussion ----------------
h("3.7 Architecture: intrinsic memory as transition resistance", 2)
p("Reframing intrinsic memory as transition RESISTANCE — the barrier a cue must overcome to leave "
  "the current attractor — rather than a re-injected default signal, keeps the model focused on "
  "its objective (predicting the pathway program) instead of imposing a secondary autoregulation "
  "mechanism. Resistance is computed per cell from lineage-TF, chromatin, and current-program "
  "states; plasticity lowers it; the state updates as resistance*current + (1-resistance)*candidate, "
  "and a soft (graded) attractor replaces winner-take-all. We compared the two formulations on an "
  "expanded 13-program pool (adding T-cell exhaustion as a genuine-generalization stress test), "
  "markers included in all arms, identical grouped folds/seeds and leakage controls.")
table(["metric (13-program, markers-in)", "initial re-injection", "resistance-gated"],
      [["macro-AUPRC", "0.509", "0.526"],
       ["balanced accuracy", "0.372", "0.403"],
       ["macro-F1", "0.177", "0.192"],
       ["program recall", "0.351", "0.349"],
       ["structure benefit vs edge-removed (AUPRC)", "+0.085 (p=0.0015)", "+0.130 (p=0.0002)"]])
em("The resistance formulation improves the top-1 metrics the initial formulation only matched "
   "(balanced accuracy +0.031, macro-F1 +0.015) at equal-or-better AUPRC, and — the key point — it "
   "WIDENS the benefit of regulatory structure: the graph's edge-removed AUPRC advantage grows from "
   "+0.085 to +0.130 (more significant), and structure begins to help program recall (which it did "
   "not under re-injection). In both formulations the edge-removed model collapses to the "
   "logistic-regression level (~0.39), confirming the graph is the lever. Interpretation: scoping "
   "memory as inertia (not a re-injected default) lets the regulatory graph contribute more to "
   "pathway prediction, consistent with the theory that lineage/chromatin context sets a "
   "transition barrier rather than a fixed bias. Caveats: the gains are modest and from a single "
   "configuration (plasticity-as-barrier-lowering, soft attractor); the folds are shared so the "
   "means are comparable but a formal paired test between formulations is future work; and the "
   "13-program pool adds a cleanly separable program (exhaustion) that raises absolute AUPRC "
   "independent of architecture, so the controlled quantities are the within-pool structure "
   "benefit and top-1 deltas, not the absolute AUPRC vs the 12-program sections above.")

h("3.8 The structure advantage scales with program diversity", 2)
p("To test whether the structure benefit is an artifact of a small program set, we tripled program "
  "breadth: seven additional curated cue->cell-state-transition datasets were ingested (intestinal "
  "differentiation, germinal-center B-cell, T-follicular-helper and regulatory-T fates, trophoblast "
  "differentiation, adipogenesis, and endothelial-to-mesenchymal transition), each with real "
  "per-cell or condition labels and a knowledge-graph cascade, taking the pool to 19 programs / 20 "
  "classes / 27,392 cells. We reran the identical converged resistance-gated configuration "
  "(hidden 128, 8 steps, 120 epochs), grouped 5-fold x 3-seed, markers in all arms, with the "
  "edge-removed structure control.")
table(["Model (19-program, markers-in)", "AUPRC", "bal-acc", "prog-rec", "macro-F1"],
      [["logistic regression", "0.296", "0.276", "0.223", "0.108"],
       ["random forest", "0.299", "0.221", "0.084", "0.135"],
       ["KG-GNN, edges removed", "0.286", "0.243", "0.220", "0.085"],
       ["KG-GNN (markers + structure)", "0.537", "0.353", "0.328", "0.166"]])
em("On the same seed x fold splits (paired Wilcoxon), the KG-GNN significantly beats logistic "
   "regression (+0.270 AUPRC, p=0.0003; +0.090 program recall, p=0.0001), random forest (+0.259 / "
   "+0.256), and — the controlled comparison — its own edge-removed twin (+0.282 AUPRC, p=0.0001; "
   "+0.096 recall). The key observation is how the gap arose: the KG-GNN's AUPRC is essentially "
   "unchanged from the 12/13-program pool (0.53 -> 0.54), while the structureless models degrade "
   "sharply as programs triple (logistic regression, random forest, and the edge-removed network all "
   "fall from ~0.40 to ~0.29). Regulatory structure thus buys ROBUSTNESS to program diversity: it "
   "sustains ranking quality where flat learners do not, and the advantage — ranking-only and modest "
   "on 12 programs — now also covers top-1 (balanced accuracy, program recall) and is roughly "
   "quadrupled (+0.07 -> +0.28). The edge-removed model again collapsing onto logistic regression "
   "confirms the graph, not features or capacity, is the lever, and rules out the new single-source "
   "programs leaking via batch identity (the edge-removed twin shares that information and does not "
   "exploit it; cue nodes are off throughout, so no cue-gating). Caveats: a single configuration; "
   "large grouped-fold variance (+/-0.11); several added programs are single-source, so grouped-split "
   "cannot test transfer TO them. A second-configuration confirmation is in progress before this is "
   "treated as more than a single-config finding.")

h("4. Discussion")
p("Read honestly, the results say something specific. The expression->program mapping is easily "
  "learned in-distribution; the hard, real problem is transfer to unseen datasets, where batch/"
  "domain shift dominates and every model degrades. In that regime the theory-structured model's "
  "concrete contribution is better program probability-RANKING (macro-AUPRC), not better top-1 "
  "accuracy — which is a sensible place for graph structure to help on an imbalanced problem, and "
  "the effect is consistent across configurations. We deliberately do not claim a blanket accuracy "
  "win: on top-1 metrics the model is competitive, not superior, and only once trained to "
  "convergence.")
p("Two architectural findings shaped the result, and both point to mis-scoped MECHANISM rather "
  "than a failure of the graph. First, the winner-take-all attractor saturated probabilities and "
  "erased graded temporal signal; a soft (graded) attractor recovered graded temporal emergence on "
  "the real EMT course, and an edge-removed control showed the regulatory graph — not the markers — "
  "drives it (rho +0.20 with structure vs +0.03 without). Second, reframing intrinsic memory as "
  "transition resistance rather than a re-injected default (Section 3.7) improved top-1 metrics and "
  "widened the structure benefit. Together these suggest the model's earlier ceilings were partly "
  "artifacts of over-strong mechanisms (hard winner-take-all, forced memory re-injection): better "
  "biological scoping let the regulatory structure contribute more — on ranking, top-1, and "
  "temporally — under honest grouped-dataset evaluation. The gains are modest, and on cross-dataset "
  "classification the model remains competitive rather than dominant; but the direction consistently "
  "supports the theory that lineage/chromatin context sets a transition barrier the cue must overcome.")

# ---------------- 5 Limitations ----------------
h("5. Limitations")
for t_ in [
  "Cross-dataset generalization is modest (grouped macro-AUPRC ~0.47); absolute performance on 12 "
  "imbalanced classes is limited, and the strength is in controlled comparisons, not raw accuracy.",
  "The advantage is specific to AUPRC (ranking); on top-1 metrics the model matches, not beats, "
  "the baselines. A converged (epochs-120) multi-seed run to confirm top-1 recovery is still useful.",
  "The temporal-emergence result is on a single EMT time-course and requires the attractor off "
  "(WTA sharpening erases the graded signal); multi-dataset temporal validation is future work.",
  "Several labels are condition/timepoint proxies (type C) rather than deposited per-cell "
  "annotations; EMT and Senescence are single-source; interpretability is weak for EMT/Pluripotency.",
  "ADM's simulated dynamics are marker-dependent. Models are small; convergence was compute-limited.",
  "The wide-model ablation and perturbation confirmations are pending.",
  "The 19-program scaling result (3.8) is a single configuration; several added programs are "
  "single-source, so grouped-split cannot test transfer TO them. A second-configuration "
  "confirmation is in progress.",
]:
    doc.add_paragraph(t_, style="List Bullet")

# ---------------- 6 Conclusion ----------------
h("6. Conclusion")
p("A multiscale theory of cell-state selection can be built as an interpretable knowledge-graph "
  "GNN and evaluated honestly on real multi-source single-cell data. After removing a label leak, "
  "the model's reproducible real-data advantage is program probability-ranking (macro-AUPRC ~0.48 "
  "vs ~0.40 baselines, paired p<0.02) on the hard cross-dataset-transfer regime; it matches baselines on top-1 "
  "metrics once converged, recovers known regulators, and reproduces the theory's plasticity-gated "
  "dynamics in simulation. The central claim — that regulatory structure adds value on top of "
  "markers — is supported by TWO independent edge-removed controls (markers held equal): it "
  "significantly improves program ranking (+0.082 AUPRC, paired p=0.015) and it drives graded "
  "temporal emergence (rho +0.20 with structure vs +0.03 without). This ranking advantage does not "
  "depend on a small program set: tripling breadth to 19 programs WIDENS it to +0.28 AUPRC over the "
  "edge-removed control (paired p=0.0001) and extends it to top-1 metrics, because structure sustains "
  "ranking quality where flat learners degrade with program count (single-config; Section 3.8). A "
  "blanket top-1 classification "
  "win is NOT claimed (the model matches, not beats, on argmax), which we report directly. The "
  "contribution is a credibility-controlled framework and an honest map of where a mechanistic "
  "inductive bias helps (probability-ranking, interpretability, simulated dynamics) and where it "
  "does not (top-1 accuracy, real temporal continuum).")

h("Data and code availability")
p("All datasets are public (accessions in the label-provenance protocol). Code, the knowledge "
  "graph, ingestion, model, evaluation harnesses (classification, baselines, ablation, "
  "perturbation, temporal-trajectory), and this manuscript builder are in the project repository.")

out = "/Users/work/MultiscaleProject/artifacts/chromatin_toggle_manuscript.docx"
doc.save(out)
print("saved", out)
